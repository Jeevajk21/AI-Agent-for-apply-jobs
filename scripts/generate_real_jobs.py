"""
Generate Word (.docx) CVs, cover letters and tracker for 20 REAL active finance jobs
found on LinkedIn (verified URLs from live search results June 2026).
"""

import os, csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = "/home/user/AI-Agent-for-apply-jobs"
CV_DIR  = os.path.join(BASE, "applications", "cvs")
CL_DIR  = os.path.join(BASE, "applications", "coverletters")
TRACKER = os.path.join(BASE, "applications", "tracker.csv")
os.makedirs(CV_DIR, exist_ok=True)
os.makedirs(CL_DIR, exist_ok=True)

BLACK = RGBColor(0, 0, 0)
TODAY = "11 June 2026"

CANDIDATE = {
    "name":     "Jeevakumar Jayachandran",
    "email":    "jeevajk2112@gmail.com",
    "phone":    "+44 7466 480480",
    "linkedin": "linkedin.com/in/jeevakumar-j",
    "location": "Lancaster, UK",
}

# ── 20 REAL jobs with verified LinkedIn URLs ──────────────────────────────────
JOBS = [
    {
        "title": "Graduate Pricing Analyst",
        "company": "Somerset Bridge Group",
        "location": "Cribbs Causeway, Bristol, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 94,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-pricing-analyst-at-somerset-bridge-group-4399764858",
        "cl_hook": "Somerset Bridge Group's position as a specialist insurance group — where pricing analytics directly drives underwriting profitability — is exactly the environment where my experience building 600+ pricing models at Accenture translates immediately into commercial impact.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Pricing Analyst",
        "company": "Houst",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 92,
        "apply_url": "https://uk.linkedin.com/jobs/view/pricing-analyst-at-houst-4421806118",
        "cl_hook": "Houst's data-driven approach to dynamic pricing in the short-term rental market — optimising rates across thousands of properties in real time — is precisely the kind of large-scale pricing challenge that builds directly on my experience engineering 600+ profitability models at Accenture.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Pricing Analyst",
        "company": "Admiral Group",
        "location": "Cardiff, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 95,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-pricing-analyst-at-admiral-group-plc-4270298136",
        "cl_hook": "Admiral Group's reputation for using data and pricing analytics as a genuine competitive advantage in the UK motor and home insurance market is exactly why I am drawn to this role — my background building 600+ pricing models at Accenture maps directly to the rigorous, model-driven pricing environment Admiral is known for.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Commercial Pricing Analyst",
        "company": "Octopus Energy",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 93,
        "apply_url": "https://uk.linkedin.com/jobs/view/commercial-pricing-analyst-at-octopus-energy-4395345171",
        "cl_hook": "Octopus Energy's mission to disrupt the energy market through smarter, fairer pricing — backed by proprietary technology and a culture of analytical rigour — makes this role a compelling fit for my commercial pricing background, where I built 600+ models underpinning multimillion-dollar contract decisions at Accenture.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Pricing Analyst",
        "company": "AXA UK",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 95,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-pricing-analyst-at-axa-uk-3566689356",
        "cl_hook": "AXA UK's Retail Pricing team — supporting New Business and Renewals analytics across Home and Motor lines — represents the ideal intersection of actuarial rigour and commercial strategy where my pricing modelling and sensitivity analysis skills from Accenture can add immediate value.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Pricing Analyst",
        "company": "AXA UK",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 93,
        "apply_url": "https://uk.linkedin.com/jobs/view/pricing-analyst-at-axa-uk-4261728359",
        "cl_hook": "AXA UK's scale across general insurance and its investment in advanced pricing capabilities creates an environment where analytical pricing professionals can drive measurable business outcomes — directly aligned with my two years building and refining 600+ pricing models at Accenture.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Commercial & Pricing Analyst",
        "company": "Ashurst",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 91,
        "apply_url": "https://uk.linkedin.com/jobs/view/commercial-pricing-analyst-at-ashurst-4292520452",
        "cl_hook": "Ashurst's position as a leading global law firm means its commercial pricing function must support complex, high-value deal structures with analytical precision — exactly the environment where my experience structuring competitive pricing and evaluating deal economics at Accenture is directly applicable.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Pricing Analyst",
        "company": "Raytheon UK",
        "location": "Glenrothes, Scotland, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 90,
        "apply_url": "https://uk.linkedin.com/jobs/view/pricing-analyst-at-raytheon-uk-4376179019",
        "cl_hook": "Raytheon UK's defence contracting environment — where pricing must reflect complex programme risks, cost structures, and government tender requirements — mirrors the multimillion-dollar bid pricing methodology I developed at Accenture, where rigorous cost modelling and scenario analysis underpinned final deal strategy.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate FP&A Analyst",
        "company": "Howden Group Holdings",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 95,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-fp-a-analyst-at-howden-4328827847",
        "cl_hook": "Howden's rapid expansion across 50+ countries and its data-driven approach to insurance analytics makes it an ideal environment for an FP&A professional who thrives at the intersection of financial modelling and strategic insight.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Finance Analyst (FP&A) – Graduate",
        "company": "OSTTRA",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 93,
        "apply_url": "https://uk.linkedin.com/jobs/view/finance-analyst-fp-a-graduate-at-osttra-4083372204",
        "cl_hook": "OSTTRA's role at the heart of global financial markets post-trade infrastructure — providing clearing, compression, and risk management solutions — creates an FP&A environment of genuine complexity and scale, where my financial modelling and variance analysis skills can contribute immediately.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Analyst – 2026 Programme",
        "company": "Baringa Partners",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 91,
        "apply_url": "https://uk.linkedin.com/jobs/view/uk-2026-graduate-analyst-at-baringa-4308629166",
        "cl_hook": "Baringa's reputation as one of the UK's leading management consulting firms — ranked in the Sunday Times Best Places to Work — combined with its focus on financial services clients, makes this graduate programme the ideal environment to apply my pricing and financial analysis expertise in a client-facing setting.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Finance Analyst Graduate Apprenticeship 2026",
        "company": "Barclays",
        "location": "Glasgow, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 93,
        "apply_url": "https://uk.linkedin.com/jobs/view/finance-analyst-graduate-apprenticeship-programme-2026-glasgow-at-barclays-4299277465",
        "cl_hook": "In my FTSE 350 equity analysis project, I flagged Barclays as undervalued on a risk-adjusted basis, with systematic beta analysis supporting a Buy/Hold recommendation — a conviction built on the same analytical rigour I would bring to Barclays' own finance function.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Graduate Analyst Programme 2026 – Management Consulting",
        "company": "CF (Corporate Finance)",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 90,
        "apply_url": "https://uk.linkedin.com/jobs/view/analyst-graduate-programme-2026-management-consulting-at-cf-4418151711",
        "cl_hook": "CF's specialist focus on corporate finance advisory — where every engagement requires combining analytical rigour with clear commercial judgement — aligns directly with my experience building financial models and structuring pricing strategies at Accenture.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Graduate Analyst – Deutsche Bank Graduate Programme 2026",
        "company": "Deutsche Bank",
        "location": "London, UK",
        "work_type": "On-site",
        "salary": "Competitive",
        "score": 92,
        "apply_url": "https://uk.linkedin.com/jobs/view/deutsche-bank-graduate-analyst-programme-investment-bank-fixed-income-currencies-london-2026-at-deutsche-bank-4295837621",
        "cl_hook": "Deutsche Bank's Fixed Income & Currencies division sits at the intersection of market analytics and quantitative finance — an environment that directly builds on my MSc Finance dissertation, where I conducted panel regression analysis on 135,000+ syndicated loan tranches to study ESG pricing dynamics in debt markets.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Graduate Investment Analyst",
        "company": "Bending Spoons",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 89,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-investment-analyst-at-bending-spoons-4290193798",
        "cl_hook": "Bending Spoons' data-driven approach to company acquisitions and portfolio management — acquiring and scaling mobile apps through rigorous financial analysis — represents an exciting application of investment analysis where my DCF modelling, CAPM expertise, and Python skills can contribute from day one.",
        "cl_proof": "investment",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Financial Analyst",
        "company": "Pictet Group",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 92,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-financial-analyst-london-at-pictet-group-4166367838",
        "cl_hook": "Pictet Group's heritage as one of Europe's premier independent wealth and asset managers — combined with its long-term investment philosophy grounded in fundamental research — mirrors the rigorous, research-driven approach I took in my MSc Finance projects, from building a full FCFF DCF model for Unilever to cross-sectional CAPM analysis across the FTSE 350.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Graduate Commercial Analyst",
        "company": "targetjobs UK",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 88,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-commercial-analyst-at-targetjobs-uk-4422075867",
        "cl_hook": "The Graduate Commercial Analyst role represents an excellent opportunity to combine financial analysis with commercial strategy — directly building on my experience at Accenture, where I analysed deal economics, built pricing models, and translated complex data into commercial recommendations for senior stakeholders.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Analyst",
        "company": "Aegis Energy",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 88,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-analyst-at-aegis-energy-4412774948",
        "cl_hook": "Aegis Energy's focus on energy markets analytics and commercial strategy — at a time when energy pricing is undergoing fundamental transformation — creates an environment where my quantitative analysis, pricing modelling, and financial forecasting skills can drive real commercial value.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Investment Analyst",
        "company": "City Grad",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 89,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-investment-analyst-at-city-grad-4172507382",
        "cl_hook": "City Grad's specialisation in placing high-potential finance graduates into leading investment roles across the City makes this an ideal platform to leverage my MSc Finance training, Bloomberg Terminal expertise, and investment analysis projects into a front-office finance career.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Graduate Technical Business Analyst (Finance) 2026",
        "company": "Acturis",
        "location": "London, UK",
        "work_type": "Hybrid",
        "salary": "Competitive",
        "score": 88,
        "apply_url": "https://uk.linkedin.com/jobs/view/graduate-technical-business-analyst-london-2026-at-acturis-limited-4375135848",
        "cl_hook": "Acturis's position as the leading software provider to the UK insurance market — combining financial technology with commercial analysis — creates a unique environment where my pricing analytics background from Accenture and strong technical skills in Python and Excel directly support their insurance industry clients.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
]

# ── CV & CL content ───────────────────────────────────────────────────────────

EDUCATION = [
    ("MSc Finance", "Lancaster University Management School, Lancaster, UK", "Oct 2025 – Aug 2026",
     "First Class expected. Modules: Advanced Investment Management, Financial Modelling & Valuation, Python for Data Analysis, Financial Databases (Bloomberg Terminal). Dissertation: The ESG Debt Premium in the Syndicated Loan Market — panel regression on 135,000+ loan tranches (DealScan, Compustat, RepRisk) in R."),
    ("Master of Commerce (M.Com), 82%", "KCS College of Arts and Science, Chennai, India", "Jun 2019 – Apr 2021", ""),
    ("Bachelor of Commerce (B.Com), 72%",  "KCS College of Arts and Science, Chennai, India", "Jun 2016 – Apr 2019", ""),
]

EXPERIENCE = [
    {
        "title": "Pricing Associate (Financial Analyst)", "company": "Accenture",
        "location": "Bengaluru, India", "dates": "Jun 2023 – Sep 2025",
        "bullets": [
            "Engineered and maintained 600+ pricing and profitability models for multimillion-dollar outsourcing contracts, producing revenue and cost forecasts that underpinned executive deal-approval decisions.",
            "Structured competitive pricing strategies that maximised long-term contract profitability by partnering with senior sales and finance stakeholders to evaluate deal economics and commercial terms.",
            "Quantified margin impacts across multiple deal scenarios by running multi-variable sensitivity analysis, isolating key cost drivers and downside risks that shaped final bid strategy.",
            "Reduced manual reporting time by 30% by designing automated performance dashboards in Excel VBA, accelerating decision-making cycles across the pricing team.",
            "Surfaced margin-improvement opportunities by conducting variance and trend analysis on large pricing datasets, distilling complex figures into clear recommendations for non-financial stakeholders.",
        ],
    },
    {
        "title": "Finance Assistant", "company": "SMS LLP (Mtandt Group)",
        "location": "Chennai, India", "dates": "Feb 2022 – May 2023",
        "bullets": [
            "Produced monthly management accounts and forecasts that supported executive decision-making, consolidating financial statements into concise reporting for non-financial leaders.",
            "Reduced potential financial losses by 12% by conducting inventory and asset-performance analysis and recommending targeted provisions to management.",
            "Tracked financial performance against annual targets by supporting the budgeting cycle and delivering monthly variance analysis.",
        ],
    },
]

PROJECTS = [
    {
        "title": "Unilever PLC: Intrinsic Valuation via FCFF DCF Model",
        "bullets": [
            "Valued Unilever's intrinsic equity by building a full FCFF discounted cash flow model, stress-testing outputs through scenario and sensitivity analysis across WACC, terminal growth, and operating-margin assumptions.",
            "Framed an equity-research-style recommendation identifying Unilever as a defensive, cash-generative business with limited near-term upside, driven by macro and discount-rate risk.",
        ],
    },
    {
        "title": "FTSE 350 Portfolio: Equity Valuation & Risk Analysis",
        "bullets": [
            "Calculated beta and risk-adjusted expected returns for 10 FTSE 350 companies applying CAPM, flagging HSBC and Barclays as undervalued (Buy/Hold) and Glencore and Unilever as overvalued (Sell).",
            "Formulated a buy/sell investment rationale grounded in systematic-risk and return-risk trade-off frameworks.",
        ],
    },
]

SKILLS = [
    ("Technical Tools", "Advanced Excel (VBA, Power Query), Bloomberg Terminal, Python (pandas, data visualisation), R (panel regression), SQL, PowerPoint"),
    ("Financial Competencies", "Equity Valuation (DCF, CAPM), Pricing & Profitability Modelling, Scenario & Sensitivity Analysis, Financial Forecasting, Management Accounting, Variance Analysis"),
    ("Certifications", "CMA Intermediate, Institute of Cost Accountants of India (ICAI)"),
    ("Interests", "Chess, Painting"),
]

PROOF = {
    "pricing": (
        "At Accenture, I engineered and maintained over 600 pricing and profitability models for "
        "multimillion-dollar outsourcing contracts, structuring competitive pricing strategies that "
        "maximised long-term contract profitability. I quantified margin impacts across multiple deal "
        "scenarios through multi-variable sensitivity analysis, and reduced manual reporting time by 30% "
        "by automating performance dashboards in Excel VBA. These experiences provide exactly the "
        "commercial pricing expertise this role demands."
    ),
    "fpa": (
        "At Accenture, I developed deep financial analysis expertise — building pricing and profitability "
        "models, conducting variance analysis on large datasets, and producing management reporting that "
        "drove executive decision-making. At SMS LLP, I prepared monthly management accounts, reduced "
        "financial losses by 12% through inventory analysis, and supported the annual budgeting cycle. "
        "Together these roles gave me the dual perspective of commercial finance and operational "
        "management accounting."
    ),
    "investment": (
        "My MSc Finance at Lancaster has provided hands-on investment analysis experience: I built a "
        "full FCFF DCF model to value Unilever's intrinsic equity and applied CAPM to evaluate "
        "risk-adjusted returns across 10 FTSE 350 companies — flagging Barclays and HSBC as undervalued "
        "on a Buy/Hold basis. I use Bloomberg Terminal daily and have developed strong Python and R "
        "skills through quantitative research, including a panel regression on 135,000+ syndicated "
        "loan tranches for my MSc dissertation."
    ),
}

SKILLS_PARA = {
    "standard": (
        "My MSc Finance at Lancaster University Management School — where I am on track for a First "
        "Class — has equipped me with advanced financial modelling, quantitative analysis, and Bloomberg "
        "Terminal skills. My CMA Intermediate qualification, combined with Python, R, SQL, and Advanced "
        "Excel capabilities, ensures I can contribute immediately to a data-driven finance team while "
        "continuing to develop professionally."
    ),
    "banking": (
        "My MSc Finance dissertation — an econometric study of the ESG debt premium across 135,000+ "
        "syndicated loan tranches using DealScan, Compustat, and RepRisk in R — demonstrates both my "
        "quantitative rigour and my understanding of debt markets and credit risk dynamics. Combined "
        "with Bloomberg Terminal proficiency, Python and SQL skills, and CMA Intermediate qualification, "
        "I bring a technically sophisticated and research-grounded finance skill set."
    ),
}

# ── Formatting helpers ────────────────────────────────────────────────────────

def narrow_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.5)

def normal_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

def font(run, size, bold=False):
    run.font.name = "Calibri"; run.font.size = Pt(size)
    run.font.bold = bold; run.font.color.rgb = BLACK

def para(doc, text="", size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
         sb=0, sa=2, ls=11.5):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(ls)
    if text:
        r = p.add_run(text); font(r, size, bold)
    return p

def section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(12)
    r = p.add_run(title.upper()); font(r, 11, True)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "000000")
    pBdr.append(bot); pPr.append(pBdr)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(11.5)
    r = p.add_run(text); font(r, 10)

def two_col(doc, left, right, lb=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(11.5)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(6.73 * 1440)))
    tabs.append(tab); pPr.append(tabs)
    r1 = p.add_run(left);           font(r1, 10, lb)
    r2 = p.add_run("\t" + right);   font(r2, 10, False)

# ── Build CV ──────────────────────────────────────────────────────────────────

def make_cv(job):
    doc = Document(); narrow_margins(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Header
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(CANDIDATE["name"].upper()); font(r, 15, True)

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.line_spacing = Pt(12)
    contacts = f"{CANDIDATE['location']}   |   {CANDIDATE['email']}   |   {CANDIDATE['phone']}   |   {CANDIDATE['linkedin']}"
    r2 = p2.add_run(contacts); font(r2, 10)

    # Education
    section_heading(doc, "Education")
    for deg, inst, dates, detail in EDUCATION:
        two_col(doc, deg, dates, lb=True)
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = Pt(11.5)
        r = p.add_run(inst); font(r, 10)
        if detail:
            pd = doc.add_paragraph(); pd.paragraph_format.space_before = Pt(0)
            pd.paragraph_format.space_after = Pt(3); pd.paragraph_format.line_spacing = Pt(11.5)
            rd = pd.add_run(detail); font(rd, 10)

    # Experience
    section_heading(doc, "Professional Experience")
    for exp in EXPERIENCE:
        two_col(doc, f"{exp['title']} | {exp['company']}, {exp['location']}", exp["dates"], lb=True)
        for b in exp["bullets"]: bullet(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Projects
    section_heading(doc, "Investment & Financial Analysis Projects")
    for proj in PROJECTS:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(proj["title"]); font(r, 10, True)
        for b in proj["bullets"]: bullet(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # Skills
    section_heading(doc, "Skills & Interests")
    for label, value in SKILLS:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = Pt(11.5)
        r1 = p.add_run(f"{label}: "); font(r1, 10, True)
        r2 = p.add_run(value);        font(r2, 10)

    slug = (job["company"] + "_" + job["title"]).replace(" ", "_").replace("/","-").replace("–","-").replace("&","and").replace("(","").replace(")","").replace(",","").replace("'","")
    path = os.path.join(CV_DIR, f"{slug}_CV.docx")
    doc.save(path)
    return path, slug

# ── Build Cover Letter ────────────────────────────────────────────────────────

def make_cl(job, slug):
    doc = Document(); normal_margins(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    for line, sb, sa, bold in [
        (CANDIDATE["name"],     0, 1, True),
        (CANDIDATE["location"], 0, 1, False),
        (CANDIDATE["phone"],    0, 1, False),
        (CANDIDATE["email"],    0, 1, False),
        (CANDIDATE["linkedin"], 0, 8, False),
        (TODAY,                 0, 8, False),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        r = p.add_run(line); font(r, 11 if bold else 10.5, bold)

    for line in ["Hiring Manager", job["company"], job["location"]]:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line); font(r, 10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)
    sr = sp.add_run(f"Re: Application for {job['title']}"); font(sr, 10.5, True)

    sal = doc.add_paragraph(); sal.paragraph_format.space_after = Pt(8)
    sr2 = sal.add_run("Dear Hiring Manager,"); font(sr2, 10.5)

    for body in [job["cl_hook"], PROOF[job["cl_proof"]], SKILLS_PARA[job["cl_skills"]]]:
        bp = doc.add_paragraph(); bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(8); bp.paragraph_format.line_spacing = Pt(13)
        br = bp.add_run(body); font(br, 10.5)

    close = (f"I would welcome the opportunity to discuss how my background can contribute to "
             f"{job['company']}'s goals. I am available for an interview at your earliest convenience "
             f"and can be reached at {CANDIDATE['email']} or {CANDIDATE['phone']}.")
    cp = doc.add_paragraph(); cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(12); cp.paragraph_format.line_spacing = Pt(13)
    cr = cp.add_run(close); font(cr, 10.5)

    yp = doc.add_paragraph(); yp.paragraph_format.space_after = Pt(24)
    yr = yp.add_run("Yours sincerely,"); font(yr, 10.5)

    np = doc.add_paragraph(); np.paragraph_format.space_after = Pt(0)
    nr = np.add_run(CANDIDATE["name"]); font(nr, 10.5, True)

    path = os.path.join(CL_DIR, f"{slug}_CoverLetter.docx")
    doc.save(path)
    return path

# ── Run ───────────────────────────────────────────────────────────────────────

print("Generating Word documents for 20 REAL jobs...\n")
rows = []

for job in JOBS:
    cv_path, slug = make_cv(job)
    cl_path = make_cl(job, slug)
    print(f"  [{job['score']}/100] {job['title']} @ {job['company']}")
    rows.append({
        "Job Title":          job["title"],
        "Company":            job["company"],
        "Location":           job["location"],
        "Job Type":           job["work_type"],
        "Score":              job["score"],
        "Date Found":         "2026-06-11",
        "Application Status": "Not Started",
        "Salary":             job["salary"],
        "Visa Sponsorship":   "Check listing",
        "Apply Link":         job["apply_url"],
        "CV File":            os.path.basename(cv_path),
        "Cover Letter File":  os.path.basename(cl_path),
    })

fields = ["Job Title","Company","Location","Job Type","Score","Date Found",
          "Application Status","Salary","Visa Sponsorship","Apply Link",
          "CV File","Cover Letter File"]
with open(TRACKER, "w", newline="") as f:
    writer = csv.DictWriter(f, fieldnames=fields)
    writer.writeheader()
    writer.writerows(rows)

print(f"\n✓ 20 CVs         → {CV_DIR}/")
print(f"✓ 20 Cover letters → {CL_DIR}/")
print(f"✓ Tracker updated  → {TRACKER}")
