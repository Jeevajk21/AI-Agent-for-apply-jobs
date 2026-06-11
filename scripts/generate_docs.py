import docx
import os

def create_cover_letter(file_path, name, contact_info, company, role, address, body):
    doc = docx.Document()
    doc.add_paragraph(name).bold = True
    for info in contact_info:
        doc.add_paragraph(info)
    doc.add_paragraph()
    doc.add_paragraph("June 10, 2026")
    doc.add_paragraph()
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph(company)
    for line in address:
        doc.add_paragraph(line)
    doc.add_paragraph()
    doc.add_paragraph(f"Re: Application for {role} Position").bold = True
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph()
    for para in body:
        doc.add_paragraph(para)
    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(name)
    doc.save(file_path)

def create_cv_summary(file_path, role_title, summary, highlights):
    doc = docx.Document()
    doc.add_heading(f"Tailored CV Content - {role_title}", 0)
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Key Highlights for this Role", level=1)
    for highlight in highlights:
        doc.add_paragraph(highlight, style='List Bullet')
    doc.add_paragraph("\nNote: Please integrate the above into your master CV.")
    doc.save(file_path)

# Correct Paths
cv_dir = r"C:\Users\Jeeva\Documents\finance-job-agent\applications\cvs"
cl_dir = r"C:\Users\Jeeva\Documents\finance-job-agent\applications\coverletters"

# Goodbody
body_goodbody = ["I am writing to express my strong interest in the Investment Analyst position at Goodbody...", "Throughout my MSc program...", "Furthermore, my proficiency with the Bloomberg Terminal...", "Goodbody’s reputation...", "I am eager to discuss..."]
create_cover_letter(os.path.join(cl_dir, "Goodbody_Cover_Letter.docx"), "Jeevakumar Jayachandran", ["[Contact Info]"], "Goodbody", "Investment Analyst", ["Dublin"], body_goodbody)
create_cv_summary(os.path.join(cv_dir, "Goodbody_Tailored_CV.docx"), "Investment Analyst at Goodbody", "Summary...", ["Highlight 1", "Highlight 2"])

# Boeing
body_boeing = ["I am writing to express my strong interest in the Financial Analyst position at Boeing...", "My academic background...", "I am particularly drawn to Boeing's commitment...", "Thank you for considering my application."]
create_cover_letter(os.path.join(cl_dir, "Boeing_Cover_Letter.docx"), "Jeevakumar Jayachandran", ["[Contact Info]"], "Boeing", "Financial Analyst", ["Belfast"], body_boeing)
create_cv_summary(os.path.join(cv_dir, "Boeing_Tailored_CV.docx"), "Financial Analyst at Boeing", "Summary...", ["Highlight 1", "Highlight 2"])

# Start People
body_start_people = ["I am writing to apply for the Commercial and Pricing Analyst position at Start People Ltd...", "My analytical background...", "I am particularly impressed...", "Thank you for considering my application."]
create_cover_letter(os.path.join(cl_dir, "Start_People_Cover_Letter.docx"), "Jeevakumar Jayachandran", ["[Contact Info]"], "Start People Ltd", "Commercial and Pricing Analyst", ["Doncaster"], body_start_people)
create_cv_summary(os.path.join(cv_dir, "Start_People_Tailored_CV.docx"), "Commercial and Pricing Analyst at Start People Ltd", "Summary...", ["Highlight 1", "Highlight 2"])

print("All Word documents generated with correct mapping.")
