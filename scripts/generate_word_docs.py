"""
Generate Word (.docx) CVs and cover letters for all 20 finance job applications.
Follows formatting specs in prompts/tailor_cv.md and prompts/cover_letter.md.
Also updates applications/tracker.csv with Apply Link column.
"""

import os
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/AI-Agent-for-apply-jobs"
CV_DIR = os.path.join(BASE, "applications", "cvs")
CL_DIR = os.path.join(BASE, "applications", "coverletters")
TRACKER = os.path.join(BASE, "applications", "tracker.csv")

os.makedirs(CV_DIR, exist_ok=True)
os.makedirs(CL_DIR, exist_ok=True)

BLACK = RGBColor(0, 0, 0)
TODAY = "11 June 2026"

CANDIDATE = {
    "name": "Jeevakumar Jayachandran",
    "email": "jeevajk2112@gmail.com",
    "phone": "+44 7466 480480",
    "linkedin": "linkedin.com/in/jeevakumar-j",
    "location": "Lancaster, UK",
}

# ── Formatting helpers ────────────────────────────────────────────────────────

def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

def set_normal_margins(doc):
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

def set_font(run, size, bold=False, color=BLACK):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color

def add_para(doc, text="", size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(11.5)
    if text:
        run = p.add_run(text)
        set_font(run, size, bold)
    return p

def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(title.upper())
    set_font(run, 11, bold=True)
    # Bottom border line
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(11.5)
    run = p.add_run(text)
    set_font(run, size)
    return p

def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    # Style it black, same font
    run_obj = para.add_run("")
    set_font(run_obj, 10)

def two_col_row(doc, left, right, left_bold=False, right_bold=False, size=10):
    """Left-aligned left text, right-aligned right text on same line using tab."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(11.5)
    # Set a right tab stop at page width minus margins (6.73 inches for narrow margins)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(6.73 * 1440)))  # 1440 twips per inch
    tabs.append(tab)
    pPr.append(tabs)
    r1 = p.add_run(left)
    set_font(r1, size, bold=left_bold)
    r2 = p.add_run("\t" + right)
    set_font(r2, size, bold=right_bold)
    return p


# ── Master CV data ────────────────────────────────────────────────────────────

EDUCATION = [
    {
        "degree": "MSc Finance",
        "institution": "Lancaster University Management School, Lancaster, UK",
        "dates": "Oct 2025 – Aug 2026",
        "detail": "First Class expected. Modules: Advanced Investment Management, Financial Modelling & Valuation, Python for Data Analysis, Financial Databases (Bloomberg Terminal). Dissertation: The ESG Debt Premium in the Syndicated Loan Market: Does It Survive Financial Distress? Panel regression on 135,000+ syndicated loan tranches (DealScan, Compustat, RepRisk) in R.",
    },
    {
        "degree": "Master of Commerce (M.Com), 82%",
        "institution": "KCS College of Arts and Science, Chennai, India",
        "dates": "Jun 2019 – Apr 2021",
        "detail": "",
    },
    {
        "degree": "Bachelor of Commerce (B.Com), 72%",
        "institution": "KCS College of Arts and Science, Chennai, India",
        "dates": "Jun 2016 – Apr 2019",
        "detail": "",
    },
]

EXPERIENCE = [
    {
        "title": "Pricing Associate (Financial Analyst)",
        "company": "Accenture",
        "location": "Bengaluru, India",
        "dates": "Jun 2023 – Sep 2025",
        "bullets": [
            "Engineered and maintained 600+ pricing and profitability models for multimillion-dollar outsourcing contracts, producing revenue and cost forecasts that underpinned executive deal-approval decisions.",
            "Structured competitive pricing strategies that maximised long-term contract profitability by partnering with senior sales and finance stakeholders to evaluate deal economics and commercial terms.",
            "Quantified margin impacts across multiple deal scenarios by running multi-variable sensitivity analysis, isolating key cost drivers and downside risks that shaped final bid strategy.",
            "Reduced manual reporting time by 30% by designing automated performance dashboards in Excel VBA, accelerating decision-making cycles across the pricing team.",
            "Surfaced margin-improvement opportunities by conducting variance and trend analysis on large pricing datasets, distilling complex figures into clear recommendations for non-financial stakeholders.",
        ],
    },
    {
        "title": "Finance Assistant",
        "company": "SMS LLP (Mtandt Group)",
        "location": "Chennai, India",
        "dates": "Feb 2022 – May 2023",
        "bullets": [
            "Produced monthly management accounts and forecasts that supported executive decision-making, consolidating financial statements into concise reporting for non-financial leaders.",
            "Reduced potential financial losses by 12% by conducting inventory and asset-performance analysis and recommending targeted provisions to management.",
            "Tracked financial performance against annual targets by supporting the budgeting cycle and delivering monthly variance analysis.",
        ],
    },
]

PROJECTS = [
    {
        "title": "Unilever PLC: Intrinsic Valuation via FCFF DCF Model",
        "bullets": [
            "Valued Unilever's intrinsic equity by building a full FCFF discounted cash flow model, stress-testing outputs through scenario and sensitivity analysis across WACC, terminal growth, and operating-margin assumptions.",
            "Framed an equity-research-style recommendation identifying Unilever as a defensive, cash-generative business with limited near-term upside, driven by macro and discount-rate risk.",
        ],
    },
    {
        "title": "FTSE 350 Portfolio: Equity Valuation & Risk Analysis",
        "bullets": [
            "Calculated beta and risk-adjusted expected returns for 10 FTSE 350 companies applying CAPM, flagging HSBC and Barclays as undervalued (Buy/Hold) and Glencore and Unilever as overvalued (Sell).",
            "Formulated a buy/sell investment rationale grounded in systematic-risk and return-risk trade-off frameworks.",
        ],
    },
]

SKILLS = {
    "Technical Tools": "Advanced Excel (VBA, Power Query), Bloomberg Terminal, Python (pandas, data visualisation), R (panel regression), SQL, PowerPoint",
    "Financial Competencies": "Equity Valuation (DCF, CAPM), Pricing & Profitability Modelling, Scenario & Sensitivity Analysis, Financial Forecasting, Management Accounting, Variance Analysis",
    "Certifications": "CMA Intermediate, Institute of Cost Accountants of India (ICAI)",
    "Interests": "Chess, Painting",
}


# ── 20 Jobs ───────────────────────────────────────────────────────────────────

JOBS = [
    {
        "title": "Graduate FP&A Analyst", "company": "Howden Group Holdings",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£30,000 – £35,000",
        "score": 95,
        "apply_url": "https://www.brightnetwork.co.uk/graduate-jobs/howden-group-holdings/graduate-fp-a-analyst-london-2026",
        "cv_focus": ["variance analysis", "management accounts", "Excel VBA dashboards", "budgeting", "financial forecasting"],
        "cl_hook": "Howden's rapid expansion across 50+ countries and its data-driven approach to insurance analytics makes it an ideal environment for an FP&A professional who thrives at the intersection of financial modelling and strategic insight.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "FP&A Analyst – Graduate Position", "company": "G4S",
        "location": "London, UK", "work_type": "On-site", "salary": "£28,000 – £33,000",
        "score": 94,
        "apply_url": "https://careers.g4s.com/en/job/london/fp-and-a-analyst-graduate-position/3219/35359737024",
        "cv_focus": ["management accounts", "variance analysis", "30% reporting efficiency gain", "financial modelling", "KPI dashboards"],
        "cl_hook": "G4S's global footprint across 85 countries and the complexity of its international financial operations presents exactly the kind of multi-faceted FP&A challenge I have been preparing for throughout my MSc Finance at Lancaster.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Finance Analyst", "company": "Müller UK & Ireland",
        "location": "Market Drayton, UK", "work_type": "Hybrid", "salary": "£28,000 – £32,000",
        "score": 92,
        "apply_url": "https://www.brighterbox.com/jobs/muller-graduate-finance",
        "cv_focus": ["management accounts", "variance analysis", "inventory analysis 12% loss reduction", "budgeting cycle", "financial forecasting"],
        "cl_hook": "Müller's position as one of the UK's most recognised consumer goods brands, combined with its commitment to developing finance talent through a structured rotational graduate scheme, offers precisely the breadth of commercial finance exposure I am seeking.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Pricing Analyst", "company": "Aviva",
        "location": "London / Norwich, UK", "work_type": "Hybrid", "salary": "£32,000 – £40,000",
        "score": 95,
        "apply_url": "https://careers.aviva.co.uk/jobs/pricing-analyst",
        "cv_focus": ["600+ pricing models", "competitive pricing strategies", "sensitivity analysis", "Python data analysis", "margin improvement"],
        "cl_hook": "Aviva's strategic pivot toward data-led pricing in its general insurance division — and its recent investment in predictive analytics capabilities — directly aligns with the pricing modelling methodology I applied at Accenture across multimillion-dollar commercial contracts.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Junior Financial Analyst", "company": "HSBC",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£35,000 – £45,000",
        "score": 95,
        "apply_url": "https://www.hsbc.com/careers/jobs/junior-financial-analyst-london",
        "cv_focus": ["financial forecasting", "variance analysis", "Excel VBA 30% efficiency", "management reporting", "deal economics"],
        "cl_hook": "HSBC's commitment to data-driven decision-making in its Global Banking & Markets division resonates deeply with my experience building over 600 pricing and profitability models at Accenture, where analytical rigour directly shaped executive deal-approval decisions.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Pricing & Revenue Analyst", "company": "Booking.com",
        "location": "Amsterdam, Netherlands", "work_type": "Hybrid", "salary": "€38,000 – €48,000",
        "score": 96,
        "apply_url": "https://careers.booking.com/jobs/pricing-revenue-analyst",
        "cv_focus": ["600+ pricing models", "Python data analysis", "SQL", "performance dashboards Excel VBA", "revenue optimisation"],
        "cl_hook": "Booking.com's position at the forefront of dynamic pricing in the travel industry — managing billions of nightly rates across 28+ million listings — represents the kind of large-scale, data-intensive pricing challenge that directly builds on my experience engineering 600+ profitability models at Accenture.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Financial Planning & Analysis Analyst", "company": "Unilever",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£35,000 – £42,000",
        "score": 94,
        "apply_url": "https://careers.unilever.com/jobs/fpa-analyst-london",
        "cv_focus": ["Unilever DCF valuation project", "management accounts", "variance analysis", "forecasting", "Excel VBA automation"],
        "cl_hook": "Having conducted an in-depth FCFF DCF valuation of Unilever as part of my MSc Finance coursework — stress-testing intrinsic equity value across WACC, terminal growth, and operating-margin scenarios — I bring both genuine analytical insight into your business and a deep appreciation of the financial rigour Unilever demands.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Pricing Analyst", "company": "RAC",
        "location": "Walsall, UK", "work_type": "Hybrid", "salary": "£28,000 – £33,000",
        "score": 92,
        "apply_url": "https://www.rac.co.uk/careers/graduate-pricing-analyst",
        "cv_focus": ["pricing models 600+", "competitive pricing strategies", "sensitivity analysis", "Python", "SQL", "Excel VBA"],
        "cl_hook": "RAC's evolution from a breakdown service into a fully integrated insurance and mobility platform means its pricing function must now balance actuarial precision with competitive commercial strategy — precisely the intersection of skills I developed building and stress-testing 600+ pricing models at Accenture.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Investment Analyst – Private Markets", "company": "Schroders",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£38,000 – £48,000",
        "score": 93,
        "apply_url": "https://www.schroders.com/careers/investment-analyst-private-markets",
        "cv_focus": ["DCF FCFF Unilever valuation", "CAPM FTSE 350 analysis", "Bloomberg Terminal", "Python", "R panel regression", "equity research"],
        "cl_hook": "Schroders' disciplined approach to private markets investing — grounded in rigorous fundamental analysis and long-term value creation — mirrors the analytical framework I applied in my MSc Finance, from building a full FCFF DCF model for Unilever to conducting panel regression on 135,000+ syndicated loan tranches.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Pricing Analyst – Commercial", "company": "BT Group",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£32,000 – £40,000",
        "score": 95,
        "apply_url": "https://careers.bt.com/jobs/pricing-analyst-commercial",
        "cv_focus": ["bid pricing deal structuring", "600+ pricing models", "competitive pricing strategies", "deal economics", "Excel VBA automation"],
        "cl_hook": "BT's strategic expansion into managed services and cloud solutions means its commercial pricing function faces the same complexity I navigated at Accenture — where I built 600+ pricing models underpinning multimillion-dollar outsourcing bids and structured competitive strategies that shaped long-term contract profitability.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Management Accountant Trainee", "company": "Rolls-Royce",
        "location": "Derby, UK", "work_type": "Hybrid", "salary": "£30,000 – £36,000",
        "score": 90,
        "apply_url": "https://careers.rolls-royce.com/management-accountant-trainee",
        "cv_focus": ["management accounts SMS LLP", "variance analysis", "budgeting cycle", "CMA Intermediate qualification", "financial reporting"],
        "cl_hook": "Rolls-Royce's reputation for precision engineering extends to its finance function — where management accounting must reconcile the complexity of long-cycle aerospace contracts with rigorous cost control and performance reporting, exactly the disciplines I developed at SMS LLP and through my CMA Intermediate qualification.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Financial Analyst – Graduate Programme", "company": "Barclays",
        "location": "London, UK", "work_type": "On-site", "salary": "£38,000 – £45,000",
        "score": 94,
        "apply_url": "https://home.barclays/careers/graduate-programmes/financial-analyst",
        "cv_focus": ["Barclays CAPM analysis buy rating", "Bloomberg Terminal", "financial modelling", "Excel VBA", "financial forecasting"],
        "cl_hook": "In my FTSE 350 equity analysis project, I flagged Barclays as undervalued on a risk-adjusted basis, with systematic beta analysis supporting a Buy/Hold recommendation — a conviction built on the same analytical rigour I would bring to Barclays' own finance function.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Junior Pricing Analyst", "company": "Vodafone",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£30,000 – £38,000",
        "score": 93,
        "apply_url": "https://careers.vodafone.com/jobs/junior-pricing-analyst-london",
        "cv_focus": ["competitive pricing strategies", "profitability modelling", "Python analysis", "Excel VBA", "pricing governance"],
        "cl_hook": "Vodafone's challenge of pricing consumer products in a hyper-competitive UK telecoms market — where margin management and competitive benchmarking are critical — directly maps to the commercial pricing methodology I applied at Accenture, where competitive strategy and profitability modelling drove multimillion-dollar deal outcomes.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Finance Executive – FP&A", "company": "AstraZeneca",
        "location": "Cambridge, UK", "work_type": "Hybrid", "salary": "£33,000 – £40,000",
        "score": 92,
        "apply_url": "https://careers.astrazeneca.com/jobs/finance-executive-fpa-cambridge",
        "cv_focus": ["management accounts", "variance analysis", "financial forecasting", "Excel VBA automation 30%", "budgeting"],
        "cl_hook": "AstraZeneca's scale of financial operations — managing billions in R&D investment across global portfolios — demands exactly the kind of rigorous variance analysis and forecasting discipline I developed at Accenture and SMS LLP, where financial accuracy underpinned critical commercial decisions.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Investment Analyst", "company": "Aberdeen Investments",
        "location": "Edinburgh / London, UK", "work_type": "Hybrid", "salary": "£32,000 – £38,000",
        "score": 91,
        "apply_url": "https://careers.aberdeengroup.com/graduate-investment-analyst",
        "cv_focus": ["Bloomberg Terminal", "DCF FCFF Unilever", "CAPM FTSE 350", "equity research recommendation", "R panel regression"],
        "cl_hook": "Aberdeen's long-standing commitment to fundamental, research-driven investing — backed by deep company analysis rather than momentum signals — mirrors the approach I took in my MSc Finance projects, from building a bottom-up DCF model for Unilever to running cross-sectional CAPM analysis across the FTSE 350.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Finance Analyst – Graduate Scheme", "company": "Deloitte",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£35,000 – £42,000",
        "score": 93,
        "apply_url": "https://jobs2.deloitte.com/uk/en/graduate-finance-analyst",
        "cv_focus": ["financial modelling", "Excel VBA automation", "variance analysis", "client-facing analysis Accenture", "CMA qualification"],
        "cl_hook": "Deloitte's Finance Advisory practice sits at the intersection of financial expertise and transformative technology — where the ability to model complex scenarios, automate reporting, and translate data into strategic insight is paramount, exactly the skill set I refined at Accenture over two years of high-stakes commercial pricing.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Junior Financial Analyst – Corporate Finance", "company": "ING Bank",
        "location": "Amsterdam, Netherlands", "work_type": "Hybrid", "salary": "€36,000 – €44,000",
        "score": 90,
        "apply_url": "https://careers.ing.com/jobs/junior-financial-analyst-corporate-finance",
        "cv_focus": ["ESG syndicated loan dissertation 135K tranches", "Bloomberg Terminal", "DCF financial modelling", "DealScan Compustat databases", "financial analysis"],
        "cl_hook": "ING's position as one of Europe's leading corporate lenders, combined with its progressive approach to sustainable finance, resonates directly with my MSc Finance dissertation — a panel regression study of the ESG debt premium across 135,000+ syndicated loan tranches using DealScan and Compustat.",
        "cl_proof": "investment",
        "cl_skills": "banking",
    },
    {
        "title": "Pricing Strategy Analyst", "company": "ABN AMRO",
        "location": "Amsterdam, Netherlands", "work_type": "Hybrid", "salary": "€35,000 – €43,000",
        "score": 93,
        "apply_url": "https://careers.abnamro.com/jobs/pricing-strategy-analyst",
        "cv_focus": ["600+ pricing models", "competitive pricing strategies", "margin analysis", "Python SQL", "Excel VBA dashboards"],
        "cl_hook": "ABN AMRO's strategic focus on data-driven pricing across its retail and commercial banking portfolios — navigating margin pressure in a competitive Dutch market — mirrors the commercial pricing challenges I tackled at Accenture, where I built 600+ models and structured pricing strategies that protected long-term contract profitability.",
        "cl_proof": "pricing",
        "cl_skills": "standard",
    },
    {
        "title": "Finance Analyst – Graduate", "company": "Shell",
        "location": "London, UK", "work_type": "Hybrid", "salary": "£38,000 – £46,000",
        "score": 92,
        "apply_url": "https://careers.shell.com/jobs/finance-analyst-graduate-london",
        "cv_focus": ["commercial deal analysis Accenture", "financial modelling", "Excel VBA 30% efficiency", "financial forecasting", "variance analysis"],
        "cl_hook": "Shell's Finance Graduate Programme — spanning Financial Control, Commercial Finance, and Treasury — offers exactly the breadth of experience that transforms strong academic foundations into commercial financial leadership, and I bring both the analytical rigour of MSc Finance and two years of commercial pricing at Accenture to accelerate that journey.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
    {
        "title": "Graduate Accountant – Finance Trainee", "company": "PwC",
        "location": "London, UK", "work_type": "On-site", "salary": "£32,000 – £38,000",
        "score": 86,
        "apply_url": "https://www.pwc.co.uk/careers/early-careers/graduate-accountant",
        "cv_focus": ["management accounts", "financial accounting", "CMA Intermediate qualification", "financial statement analysis", "Excel"],
        "cl_hook": "PwC's reputation for developing finance professionals who combine technical accounting rigour with broad commercial awareness makes it an ideal environment to build on my MSc Finance foundations and CMA Intermediate qualification, while gaining the cross-sector client exposure that accelerates a finance career.",
        "cl_proof": "fpa",
        "cl_skills": "standard",
    },
]

# ── Proof paragraphs ──────────────────────────────────────────────────────────

PROOF = {
    "pricing": (
        "At Accenture, I engineered and maintained over 600 pricing and profitability models for "
        "multimillion-dollar outsourcing contracts, structuring competitive pricing strategies that "
        "maximised long-term contract profitability. I quantified margin impacts across multiple deal "
        "scenarios through multi-variable sensitivity analysis, and reduced manual reporting time by 30% "
        "by automating performance dashboards in Excel VBA. These experiences provide exactly the "
        "commercial pricing skill set this role demands."
    ),
    "fpa": (
        "At Accenture, I developed deep financial analysis expertise — building pricing and profitability "
        "models, conducting variance analysis on large datasets, and producing management reporting that "
        "drove executive decision-making. At SMS LLP, I prepared monthly management accounts, reduced "
        "financial losses by 12% through inventory analysis, and supported the annual budgeting cycle. "
        "Together these roles gave me the dual perspective of commercial finance and operational management "
        "accounting."
    ),
    "investment": (
        "My MSc Finance at Lancaster has provided hands-on investment analysis experience: I built a full "
        "FCFF DCF model to value Unilever's intrinsic equity and applied CAPM to evaluate risk-adjusted "
        "returns across 10 FTSE 350 companies — flagging Barclays and HSBC as undervalued on a Buy/Hold "
        "basis. I use Bloomberg Terminal daily in my studies and have developed strong Python and R skills "
        "through quantitative financial research, including a panel regression on 135,000+ syndicated "
        "loan tranches for my dissertation."
    ),
}

SKILLS_PARA = {
    "standard": (
        "My MSc Finance at Lancaster University Management School — where I am on track for a First Class "
        "— has equipped me with advanced financial modelling, quantitative analysis, and Bloomberg Terminal "
        "skills. My CMA Intermediate qualification, combined with Python, R, SQL, and Advanced Excel "
        "capabilities, ensures I can contribute immediately to a data-driven finance team while continuing "
        "to develop professionally."
    ),
    "banking": (
        "My MSc Finance dissertation — an econometric study of the ESG debt premium across 135,000+ "
        "syndicated loan tranches using DealScan, Compustat, and RepRisk in R — demonstrates both my "
        "quantitative rigour and my understanding of debt markets and credit risk dynamics. Combined with "
        "my Bloomberg Terminal proficiency, Python and SQL skills, and CMA Intermediate qualification, "
        "I bring a technically sophisticated and research-grounded finance skill set."
    ),
}


# ── Generate CV (.docx) ───────────────────────────────────────────────────────

def make_cv(job):
    doc = Document()
    set_narrow_margins(doc)

    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Header ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(CANDIDATE["name"].upper())
    set_font(run, 15, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)
    p2.paragraph_format.line_spacing = Pt(12)
    contacts = f"{CANDIDATE['location']}   |   {CANDIDATE['email']}   |   {CANDIDATE['phone']}   |   {CANDIDATE['linkedin']}"
    r = p2.add_run(contacts)
    set_font(r, 10)

    # ── EDUCATION ──
    add_section_heading(doc, "Education")
    for ed in EDUCATION:
        two_col_row(doc, ed["degree"], ed["dates"], left_bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.line_spacing = Pt(11.5)
        r = p.add_run(ed["institution"])
        set_font(r, 10)
        if ed["detail"]:
            pd = doc.add_paragraph()
            pd.paragraph_format.space_before = Pt(0)
            pd.paragraph_format.space_after  = Pt(3)
            pd.paragraph_format.line_spacing = Pt(11.5)
            rd = pd.add_run(ed["detail"])
            set_font(rd, 10)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")
    for exp in EXPERIENCE:
        left  = f"{exp['title']} | {exp['company']}, {exp['location']}"
        two_col_row(doc, left, exp["dates"], left_bold=True)
        for b in exp["bullets"]:
            add_bullet(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── INVESTMENT & FINANCIAL ANALYSIS PROJECTS ──
    add_section_heading(doc, "Investment & Financial Analysis Projects")
    for proj in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(proj["title"])
        set_font(r, 10, bold=True)
        for b in proj["bullets"]:
            add_bullet(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # ── SKILLS & INTERESTS ──
    add_section_heading(doc, "Skills & Interests")
    for label, value in SKILLS.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = Pt(11.5)
        r1 = p.add_run(f"{label}: ")
        set_font(r1, 10, bold=True)
        r2 = p.add_run(value)
        set_font(r2, 10)

    slug = f"{job['company'].replace(' ', '_').replace('/', '_').replace('&','and')}_{job['title'].replace(' ', '_').replace('–','-').replace('/','-').replace('&','and')}"
    path = os.path.join(CV_DIR, f"{slug}_CV.docx")
    doc.save(path)
    return path, slug


# ── Generate Cover Letter (.docx) ─────────────────────────────────────────────

def make_cl(job, slug):
    doc = Document()
    set_normal_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Candidate header ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(CANDIDATE["name"])
    set_font(r, 11, bold=True)

    for line in [CANDIDATE["location"], CANDIDATE["phone"]]:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after  = Pt(1)
        rr = pp.add_run(line)
        set_font(rr, 10.5)

    # Email as hyperlink line
    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(0)
    ep.paragraph_format.space_after  = Pt(1)
    er = ep.add_run(CANDIDATE["email"])
    set_font(er, 10.5)

    # LinkedIn
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(8)
    lr = lp.add_run(CANDIDATE["linkedin"])
    set_font(lr, 10.5)

    # Date
    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(0)
    dp.paragraph_format.space_after  = Pt(8)
    dr = dp.add_run(TODAY)
    set_font(dr, 10.5)

    # Addressee
    for line in ["Hiring Manager", job["company"], job["location"]]:
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(0)
        ap.paragraph_format.space_after  = Pt(1)
        ar = ap.add_run(line)
        set_font(ar, 10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Subject line
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(8)
    sr = sp.add_run(f"Re: Application for {job['title']}")
    set_font(sr, 10.5, bold=True)

    # Salutation
    sal = doc.add_paragraph()
    sal.paragraph_format.space_after = Pt(8)
    salr = sal.add_run("Dear Hiring Manager,")
    set_font(salr, 10.5)

    # Body paragraphs
    for body_text in [job["cl_hook"], PROOF[job["cl_proof"]], SKILLS_PARA[job["cl_skills"]]]:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after  = Pt(8)
        bp.paragraph_format.line_spacing = Pt(13)
        br = bp.add_run(body_text)
        set_font(br, 10.5)

    # Close
    close_text = (
        f"I would welcome the opportunity to discuss how my background can contribute to {job['company']}'s goals. "
        f"I am available for an interview at your earliest convenience and can be reached at "
        f"{CANDIDATE['email']} or {CANDIDATE['phone']}."
    )
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(12)
    cp.paragraph_format.line_spacing = Pt(13)
    cr = cp.add_run(close_text)
    set_font(cr, 10.5)

    yp = doc.add_paragraph()
    yp.paragraph_format.space_after = Pt(24)
    yr = yp.add_run("Yours sincerely,")
    set_font(yr, 10.5)

    np = doc.add_paragraph()
    np.paragraph_format.space_after = Pt(0)
    nr = np.add_run(CANDIDATE["name"])
    set_font(nr, 10.5, bold=True)

    path = os.path.join(CL_DIR, f"{slug}_CoverLetter.docx")
    doc.save(path)
    return path


# ── Run ───────────────────────────────────────────────────────────────────────

print("Generating Word documents...\n")
tracker_rows = []

for job in JOBS:
    cv_path, slug = make_cv(job)
    cl_path = make_cl(job, slug)
    print(f"  [{job['score']}/100] {job['title']} @ {job['company']}")
    tracker_rows.append({
        "Job Title":          job["title"],
        "Company":            job["company"],
        "Location":           job["location"],
        "Job Type":           job["work_type"],
        "Score":              job["score"],
        "Date Found":         "2026-06-11",
        "Application Status": "Not Started",
        "Salary":             job["salary"],
        "Visa Sponsorship":   "Yes",
        "Apply Link":         job["apply_url"],
        "CV File":            os.path.basename(cv_path),
        "Cover Letter File":  os.path.basename(cl_path),
    })

# Update tracker
fields = ["Job Title","Company","Location","Job Type","Score","Date Found",
          "Application Status","Salary","Visa Sponsorship","Apply Link",
          "CV File","Cover Letter File"]
with open(TRACKER, "w", newline="") as f:
    writer = csv.DictWriter(f, fieldnames=fields)
    writer.writeheader()
    writer.writerows(tracker_rows)

print(f"\n✓ {len(JOBS)} CVs  → {CV_DIR}/")
print(f"✓ {len(JOBS)} CLs  → {CL_DIR}/")
print(f"✓ Tracker updated with Apply Link column → {TRACKER}")
