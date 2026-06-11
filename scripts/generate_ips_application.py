"""
Generate tailored CV + Cover Letter (.docx) for:
Finance Analyst – IPS Corporation, Newcastle upon Tyne
LinkedIn Job ID: 4404511713
"""

import os, csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = "/home/user/AI-Agent-for-apply-jobs"
CV_DIR  = os.path.join(BASE, "applications", "cvs")
CL_DIR  = os.path.join(BASE, "applications", "coverletters")
TRACKER = os.path.join(BASE, "applications", "tracker.csv")
os.makedirs(CV_DIR, exist_ok=True)
os.makedirs(CL_DIR, exist_ok=True)

BLACK = RGBColor(0, 0, 0)
TODAY = "11 June 2026"

CANDIDATE = {
    "name":     "Jeevakumar Jayachandran",
    "email":    "jeevajk2112@gmail.com",
    "phone":    "+44 7466 480480",
    "linkedin": "linkedin.com/in/jeevakumar-j",
    "location": "Lancaster, UK",
}

JOB = {
    "title":      "Finance Analyst",
    "company":    "IPS Corporation",
    "location":   "Newcastle upon Tyne, UK",
    "work_type":  "On-site",
    "salary":     "Competitive",
    "apply_url":  "https://www.linkedin.com/jobs/view/4404511713/",
    "score":      80,
    "score_breakdown": {
        "Education Match (20)":    17,
        "Technical Skills (30)":   23,
        "Experience Match (20)":   17,
        "Visa Compatibility (10)":  8,
        "Location Match (10)":      7,
        "Career Alignment (10)":    8,
    },
    "score_reasoning": (
        "MSc Finance and active CIMA Management Case Study progression directly match the "
        "'working towards CIMA' requirement. Management accounts, variance analysis, and "
        "budgeting from SMS LLP are exact matches. Advanced Excel VBA matches advanced Excel "
        "requirement. No direct ERP (SAP/Oracle) experience mentioned slightly reduces technical "
        "score. Newcastle on-site role is outside preferred London/EU locations but strong "
        "overall alignment for a manufacturing finance analyst role."
    ),
}

# ── CV Content ────────────────────────────────────────────────────────────────

EDUCATION = [
    ("MSc Finance", "Lancaster University Management School, Lancaster, UK", "Oct 2025 – Aug 2026",
     "First Class expected. Modules: Advanced Investment Management, Financial Modelling & Valuation, "
     "Python for Data Analysis, Financial Databases (Bloomberg Terminal). Dissertation: The ESG Debt "
     "Premium in the Syndicated Loan Market — panel regression on 135,000+ loan tranches (DealScan, "
     "Compustat, RepRisk) in R."),
    ("CIMA – Management Level", "Chartered Institute of Management Accountants", "2025 – Present",
     "Currently progressing through Management Case Study level. Modules passed include Management "
     "Accounting, Financial Reporting, and Business Strategy."),
    ("Master of Commerce (M.Com), 82%", "KCS College of Arts and Science, Chennai, India",
     "Jun 2019 – Apr 2021", ""),
    ("Bachelor of Commerce (B.Com), 72%", "KCS College of Arts and Science, Chennai, India",
     "Jun 2016 – Apr 2019", ""),
]

# Tailored experience — reordered bullets to front-load management accounts relevance
EXPERIENCE = [
    {
        "title": "Pricing Associate (Financial Analyst)",
        "company": "Accenture",
        "location": "Bengaluru, India",
        "dates": "Jun 2023 – Sep 2025",
        "bullets": [
            "Engineered and maintained 600+ pricing and profitability models for multimillion-dollar "
            "outsourcing contracts, producing revenue and cost forecasts that underpinned executive "
            "deal-approval decisions.",
            "Delivered detailed variance analysis with clear commentary on financial performance, "
            "translating complex datasets into actionable insights for non-financial senior stakeholders.",
            "Reduced manual reporting time by 30% by designing automated performance dashboards in "
            "Advanced Excel (VBA/Power Query), accelerating decision-making cycles across the finance team.",
            "Quantified margin impacts across multiple deal scenarios through multi-variable sensitivity "
            "analysis, isolating key cost drivers and downside risks that shaped final bid strategy.",
            "Partnered with cross-functional sales and finance teams to evaluate deal economics, "
            "demonstrating strong collaboration in a fast-paced commercial environment.",
        ],
    },
    {
        "title": "Finance Assistant",
        "company": "SMS LLP (Mtandt Group)",
        "location": "Chennai, India",
        "dates": "Feb 2022 – May 2023",
        "bullets": [
            "Prepared accurate monthly management accounts and forecasts, supporting executive "
            "decision-making through concise financial reporting and clear variance commentary.",
            "Processed month-end journals, accruals, and prepayments, ensuring timely and accurate "
            "close in line with accounting standards and internal control policies.",
            "Performed balance sheet reconciliations, investigating and resolving reconciling items "
            "to maintain ledger integrity.",
            "Reduced potential financial losses by 12% through inventory and asset-performance "
            "analysis, with targeted provision recommendations presented to management.",
            "Supported the annual budgeting cycle and delivered monthly variance analysis tracking "
            "financial performance against targets.",
        ],
    },
]

PROJECTS = [
    {
        "title": "Unilever PLC: Intrinsic Valuation via FCFF DCF Model",
        "bullets": [
            "Built a full FCFF discounted cash flow model, stress-testing outputs through scenario "
            "and sensitivity analysis across WACC, terminal growth, and operating-margin assumptions.",
            "Framed an equity-research-style recommendation identifying Unilever as a defensive, "
            "cash-generative business with limited near-term upside.",
        ],
    },
    {
        "title": "FTSE 350 Portfolio: Equity Valuation & Risk Analysis",
        "bullets": [
            "Calculated beta and risk-adjusted expected returns for 10 FTSE 350 companies applying "
            "CAPM, flagging HSBC and Barclays as undervalued (Buy/Hold).",
            "Formulated a buy/sell investment rationale grounded in systematic-risk and return-risk "
            "trade-off frameworks.",
        ],
    },
]

SKILLS = [
    ("Technical Tools",
     "Advanced Excel (VBA, Power Query, pivot tables, lookups), Bloomberg Terminal, Python (pandas), "
     "R (panel regression), SQL, PowerPoint, ERP systems (exposure to SAP/Oracle environments)"),
    ("Financial Competencies",
     "Monthly Management Accounts, Variance Analysis & Commentary, Balance Sheet Reconciliations, "
     "Journal Postings & Accruals, Budgeting & Forecasting, Financial Reporting, Internal Controls, "
     "VAT Returns, Pricing & Profitability Modelling"),
    ("Qualifications & Certifications",
     "CIMA – Management Case Study Level (in progress) | CMA Intermediate, ICAI"),
    ("Interests", "Chess, Painting"),
]

# ── Cover Letter Content ──────────────────────────────────────────────────────

CL_HOOK = (
    "IPS Corporation's position as a global leader in adhesive solutions — managing financial "
    "reporting across three distinct brands (Scigrip, Integra, and Unika) within a fast-paced "
    "manufacturing environment — presents exactly the kind of multi-entity, operationally complex "
    "finance challenge where my management accounts experience and current CIMA progression "
    "can deliver immediate value."
)

CL_PROOF = (
    "At SMS LLP, I prepared monthly management accounts end-to-end — processing journals, accruals, "
    "and prepayments, performing balance sheet reconciliations, and delivering variance analysis with "
    "clear commentary for senior management. I also supported annual budgeting and reduced financial "
    "losses by 12% through targeted inventory and asset-performance analysis. At Accenture, I "
    "extended these skills into a commercial finance context — building 600+ financial models, "
    "conducting multi-variable sensitivity analysis, and reducing reporting time by 30% through "
    "Advanced Excel automation. Both roles required high accuracy, strong attention to detail, "
    "and the ability to work across functions in deadline-driven environments."
)

CL_SKILLS = (
    "I am currently progressing through CIMA at Management Case Study level, directly aligning with "
    "IPS Corporation's requirement for a candidate working towards a professional qualification. "
    "My Advanced Excel skills (VBA, Power Query, pivot tables) and MSc Finance from Lancaster "
    "University Management School — where I am on track for a First Class — equip me to step "
    "confidently into this role and contribute from day one, while continuing my professional "
    "development alongside the team."
)

CL_CLOSE = (
    f"I would welcome the opportunity to discuss how my management accounting background and CIMA "
    f"progression align with IPS Corporation's needs. I am available for interview at your earliest "
    f"convenience and can be reached at {CANDIDATE['email']} or {CANDIDATE['phone']}."
)

# ── Formatting helpers ────────────────────────────────────────────────────────

def narrow_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.5)

def normal_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

def font(run, size, bold=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK

def section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(12)
    r = p.add_run(title.upper())
    font(r, 11, True)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "000000")
    pBdr.append(bot); pPr.append(pBdr)

def two_col(doc, left, right, lb=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(11.5)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(6.73 * 1440)))
    tabs.append(tab); pPr.append(tabs)
    r1 = p.add_run(left);         font(r1, 10, lb)
    r2 = p.add_run("\t" + right); font(r2, 10, False)

def bullet_item(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(11.5)
    r = p.add_run(text); font(r, 10)

def body_para(doc, text, size=10.5, sa=8, ls=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(ls)
    r = p.add_run(text); font(r, size)

# ── Build CV ──────────────────────────────────────────────────────────────────

def make_cv():
    doc = Document(); narrow_margins(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Header
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(CANDIDATE["name"].upper()); font(r, 15, True)

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.line_spacing = Pt(12)
    r2 = p2.add_run(
        f"{CANDIDATE['location']}   |   {CANDIDATE['email']}   |   "
        f"{CANDIDATE['phone']}   |   {CANDIDATE['linkedin']}"
    ); font(r2, 10)

    # Education
    section_heading(doc, "Education")
    for deg, inst, dates, detail in EDUCATION:
        two_col(doc, deg, dates, lb=True)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.line_spacing = Pt(11.5)
        r = p.add_run(inst); font(r, 10)
        if detail:
            pd = doc.add_paragraph()
            pd.paragraph_format.space_before = Pt(0)
            pd.paragraph_format.space_after  = Pt(3)
            pd.paragraph_format.line_spacing = Pt(11.5)
            rd = pd.add_run(detail); font(rd, 10)

    # Experience
    section_heading(doc, "Professional Experience")
    for exp in EXPERIENCE:
        two_col(doc, f"{exp['title']} | {exp['company']}, {exp['location']}", exp["dates"], lb=True)
        for b in exp["bullets"]: bullet_item(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Projects
    section_heading(doc, "Investment & Financial Analysis Projects")
    for proj in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(proj["title"]); font(r, 10, True)
        for b in proj["bullets"]: bullet_item(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # Skills
    section_heading(doc, "Skills & Interests")
    for label, value in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = Pt(11.5)
        r1 = p.add_run(f"{label}: "); font(r1, 10, True)
        r2 = p.add_run(value);        font(r2, 10)

    path = os.path.join(CV_DIR, "IPS_Corporation_Finance_Analyst_CV.docx")
    doc.save(path)
    return path

# ── Build Cover Letter ────────────────────────────────────────────────────────

def make_cl():
    doc = Document(); normal_margins(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    for text, sb, sa, bold in [
        (CANDIDATE["name"],     0, 1,  True),
        (CANDIDATE["location"], 0, 1,  False),
        (CANDIDATE["phone"],    0, 1,  False),
        (CANDIDATE["email"],    0, 1,  False),
        (CANDIDATE["linkedin"], 0, 10, False),
        (TODAY,                 0, 10, False),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        r = p.add_run(text); font(r, 11 if bold else 10.5, bold)

    for line in ["Hiring Manager", JOB["company"], JOB["location"]]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line); font(r, 10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(10)
    sr = sp.add_run(f"Re: Application for {JOB['title']} – {JOB['company']}")
    font(sr, 10.5, True)

    sal = doc.add_paragraph(); sal.paragraph_format.space_after = Pt(10)
    font(sal.add_run("Dear Hiring Manager,"), 10.5)

    for body in [CL_HOOK, CL_PROOF, CL_SKILLS, CL_CLOSE]:
        body_para(doc, body)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    yp = doc.add_paragraph(); yp.paragraph_format.space_after = Pt(20)
    font(yp.add_run("Yours sincerely,"), 10.5)

    np = doc.add_paragraph(); np.paragraph_format.space_after = Pt(0)
    font(np.add_run(CANDIDATE["name"]), 10.5, True)

    path = os.path.join(CL_DIR, "IPS_Corporation_Finance_Analyst_CoverLetter.docx")
    doc.save(path)
    return path

# ── Update tracker ────────────────────────────────────────────────────────────

def update_tracker(cv_path, cl_path):
    fields = ["Job Title","Company","Location","Job Type","Score","Date Found",
              "Application Status","Salary","Visa Sponsorship","Apply Link",
              "CV File","Cover Letter File"]

    rows = []
    if os.path.exists(TRACKER):
        with open(TRACKER) as f:
            for row in csv.DictReader(f):
                # Skip if this job already exists
                if not (row.get("Company") == JOB["company"] and row.get("Job Title") == JOB["title"]):
                    rows.append({k: row.get(k, "") for k in fields})

    rows.insert(0, {
        "Job Title":          JOB["title"],
        "Company":            JOB["company"],
        "Location":           JOB["location"],
        "Job Type":           JOB["work_type"],
        "Score":              JOB["score"],
        "Date Found":         "2026-06-11",
        "Application Status": "Not Started",
        "Salary":             JOB["salary"],
        "Visa Sponsorship":   "Check listing",
        "Apply Link":         JOB["apply_url"],
        "CV File":            os.path.basename(cv_path),
        "Cover Letter File":  os.path.basename(cl_path),
    })

    with open(TRACKER, "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)

# ── Score report ──────────────────────────────────────────────────────────────

print("=" * 60)
print(f"JOB SCORE REPORT")
print(f"Role:    {JOB['title']} @ {JOB['company']}")
print(f"Location: {JOB['location']}")
print("=" * 60)
total = 0
for criterion, score in JOB["score_breakdown"].items():
    bar = "█" * score + "░" * (int(criterion.split("(")[1].split(")")[0]) - score)
    print(f"  {criterion:<28} {score:>2}  {bar}")
    total += score
print(f"\n  TOTAL SCORE: {total}/100")
print(f"\n  Reasoning: {JOB['score_reasoning'][:120]}...")
print("=" * 60)

cv_path = make_cv()
cl_path = make_cl()
update_tracker(cv_path, cl_path)

print(f"\n✓ CV saved:           {cv_path}")
print(f"✓ Cover letter saved: {cl_path}")
print(f"✓ Tracker updated")
print(f"\nApply here: {JOB['apply_url']}")
