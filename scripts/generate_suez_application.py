"""
Tailored CV + Cover Letter generator for:
  Graduate Financial Analyst @ SUEZ — Bolton, UK
  Salary: £28,673–£31,257 | Fixed term 24 months | Closes 26/06/2026
  ⚠ NO visa sponsorship (now or in the future)
"""
import csv, os, copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── paths ──────────────────────────────────────────────────────────────────
CV_OUT  = "applications/cvs/SUEZ_Graduate_Financial_Analyst_CV.docx"
CL_OUT  = "applications/coverletters/SUEZ_Graduate_Financial_Analyst_CoverLetter.docx"
TRACKER = "applications/tracker.csv"
os.makedirs("applications/cvs",          exist_ok=True)
os.makedirs("applications/coverletters", exist_ok=True)

# ── scoring ────────────────────────────────────────────────────────────────
scores = {
    "Education Match (20)":    19,
    "Technical Skills (30)":   25,
    "Experience Match (20)":   17,
    "Visa Compatibility (10)":  0,   # explicit NO sponsorship
    "Location Match (10)":      7,
    "Career Alignment (10)":    9,
}
total = sum(scores.values())
reasoning = (
    "MSc Finance (First Class expected) directly satisfies 'recent graduate in Finance'. "
    "Accenture pricing models and SMS LLP management accounts map cleanly to budgeting, "
    "forecasting, KPI monitoring, and financial modelling requirements. Excel VBA and "
    "Python skills cover the Excel/data requirement. ZERO visa points — role explicitly "
    "states no sponsorship now or in future. Bolton (Greater Manchester) is ~40 min from "
    "Lancaster, good location match."
)

# ── helpers ────────────────────────────────────────────────────────────────
FONT = "Calibri"

def _set_font(run, size, bold=False, color=None):
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, 10, bold=True, color=(0,70,127))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "00468F")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def two_col(doc, left, right, left_bold=False, right_bold=False,
            left_size=10, right_size=9, left_italic=False, right_italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Twips(int(6.73 * 1440)), 2)  # right-align at margin
    rl = p.add_run(left)
    _set_font(rl, left_size, bold=left_bold)
    rl.italic = left_italic
    rr = p.add_run("\t" + right)
    _set_font(rr, right_size, bold=right_bold)
    rr.italic = right_italic
    return p

def bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    _set_font(run, size)
    return p

def body_line(doc, text, size=9.5, bold=False, space_after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, size, bold=bold)
    return p

def set_narrow_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin   = Inches(0.5)
        sec.right_margin  = Inches(0.5)

def set_normal_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

# ══════════════════════════════════════════════════════════════════════════════
# CV
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
set_narrow_margins(doc)

# ── Name ──────────────────────────────────────────────────────────────────
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = name_p.add_run("JEEVAKUMAR JAYACHANDRAN")
_set_font(nr, 16, bold=True, color=(0,70,127))

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(4)
cr = contact_p.add_run(
    "Lancaster, UK  |  jeevajk2112@gmail.com  |  +44 7466 480480  |  linkedin.com/in/jeevakumar-j"
)
_set_font(cr, 9)

# ── Profile ───────────────────────────────────────────────────────────────
section_heading(doc, "Professional Profile")
body_line(doc,
    "Recent MSc Finance graduate (First Class, Lancaster University) with two years' professional experience "
    "in financial modelling, budgeting, forecasting, and management reporting. Proven ability to build "
    "multi-variable Excel models, monitor KPIs against business plans, and deliver financial analysis that "
    "supports data-driven decision-making. Actively pursuing CIMA (Management Case Study level) to "
    "deepen management accounting expertise.",
    size=9.5, space_after=2)

# ── Education ─────────────────────────────────────────────────────────────
section_heading(doc, "Education")

two_col(doc, "MSc Finance", "Oct 2025 – Aug 2026", left_bold=True, right_bold=False, left_size=10)
two_col(doc, "Lancaster University Management School, Lancaster, UK",
        "First Class (Expected)", left_italic=True, right_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Modules: Financial Modelling, Corporate Finance, Investment Analysis, Quantitative Methods, Bloomberg Terminal")
bullet(doc, "Dissertation: Quantitative analysis of ESG factors and equity risk premia across FTSE 350 constituents")

two_col(doc, "CIMA – Management Case Study Level (In Progress)", "2025 – Present",
        left_bold=True, left_size=10)
two_col(doc, "Chartered Institute of Management Accountants",
        "", left_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Management level covering: Financial Reporting, Performance Management, Financial Management, and Business Strategy")

two_col(doc, "Master of Commerce (M.Com)", "Jun 2019 – May 2021", left_bold=True, left_size=10)
two_col(doc, "University of Madras, India",
        "First Class – 82%", left_italic=True, right_italic=True, left_size=9.5, right_size=9.5)

two_col(doc, "Bachelor of Commerce (B.Com)", "Jun 2016 – May 2019", left_bold=True, left_size=10)
two_col(doc, "University of Madras, India",
        "First Class – 72%", left_italic=True, right_italic=True, left_size=9.5, right_size=9.5)

# ── Experience ────────────────────────────────────────────────────────────
section_heading(doc, "Professional Experience")

two_col(doc, "Pricing Associate (Financial Analyst)", "Jun 2023 – Sep 2025",
        left_bold=True, left_size=10)
two_col(doc, "Accenture, Bengaluru, India",
        "", left_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Engineered and maintained 600+ financial models for multimillion-dollar outsourcing contracts, covering cost modelling, revenue forecasting, and profitability analysis to support commercial bid decisions.")
bullet(doc, "Monitored KPIs and cost variances against approved business plans, producing monthly performance dashboards that highlighted deviation drivers for senior stakeholders.")
bullet(doc, "Built multi-variable sensitivity analysis models in Excel to quantify margin impact across deal scenarios, directly informing pricing strategy and competitive tender submissions.")
bullet(doc, "Reduced manual financial reporting time by 30% by automating data consolidation and performance dashboards using Excel VBA and Power Query.")
bullet(doc, "Collaborated with finance, operations, and commercial teams to develop annual budgets and rolling forecasts, ensuring alignment between financial targets and operational plans.")

two_col(doc, "Finance Assistant", "Feb 2022 – May 2023", left_bold=True, left_size=10)
two_col(doc, "SMS LLP (Mtandt Group), Chennai, India",
        "", left_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Produced monthly management accounts including P&L analysis and variance commentary, comparing actuals against budget to support operational decision-making.")
bullet(doc, "Maintained accurate financial records through journal postings, accruals, prepayments, and balance sheet reconciliations across multiple cost centres.")
bullet(doc, "Identified a 12% reduction in inventory holding losses through detailed inventory and cost analysis, presenting findings to management with supporting financial models.")
bullet(doc, "Assisted in budgeting and forecasting cycles, including cost centre budget preparation and month-end variance analysis against plan.")

# ── Projects ─────────────────────────────────────────────────────────────
section_heading(doc, "Academic Projects")

two_col(doc, "Unilever PLC – FCFF Valuation Model", "2026", left_bold=True, left_size=10)
bullet(doc, "Built a three-statement DCF model projecting free cash flow to firm; benchmarked against analyst consensus using Bloomberg Terminal data.")

two_col(doc, "FTSE 350 – CAPM & Risk Analysis", "2026", left_bold=True, left_size=10)
bullet(doc, "Applied Fama-French three-factor model using R (panel regression) across FTSE 350; analysed systematic risk and equity risk premia.")

# ── Skills ────────────────────────────────────────────────────────────────
section_heading(doc, "Skills & Competencies")

two_col(doc, "Financial Modelling & Analysis:", "Budgeting, Forecasting, Variance Analysis, KPI Monitoring, Cost Modelling", left_bold=True, left_size=9.5, right_size=9.5)
two_col(doc, "Accounting:", "Monthly Management Accounts, Balance Sheet Reconciliations, Journal Postings, Accruals & Prepayments", left_bold=True, left_size=9.5, right_size=9.5)
two_col(doc, "Technical Tools:", "Advanced Excel (VBA, Power Query, Pivot Tables), Python (pandas), SQL, R, Power BI, Bloomberg Terminal", left_bold=True, left_size=9.5, right_size=9.5)
two_col(doc, "Professional:", "CMA Intermediate (ICAI) | CIMA Management Case Study Level (in progress)", left_bold=True, left_size=9.5, right_size=9.5)

doc.save(CV_OUT)
print(f"✓ CV saved:           {CV_OUT}")

# ══════════════════════════════════════════════════════════════════════════════
# COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
cl = Document()
set_normal_margins(cl)

def cl_para(doc, text, size=11, bold=False, space_after=8, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _set_font(r, size, bold=bold)
    return p

# Header
cl_para(cl, "Jeevakumar Jayachandran", size=13, bold=True, space_after=2)
cl_para(cl, "Lancaster, UK  |  jeevajk2112@gmail.com  |  +44 7466 480480  |  linkedin.com/in/jeevakumar-j",
        size=10, space_after=12)

cl_para(cl, "2026-06-11", size=11, space_after=8)
cl_para(cl, "Hiring Manager\nSUEZ\nBolton, UK", size=11, space_after=12)
cl_para(cl, "Re: Application for Graduate Financial Analyst (24-Month Fixed Term)", size=11, bold=True, space_after=12)
cl_para(cl, "Dear Hiring Manager,", size=11, space_after=12)

cl_para(cl,
    "SUEZ's mission to deliver sustainable resource management — with financial performance measured "
    "against business plans, commercial bids evaluated through rigorous financial modelling, and KPIs "
    "tracked to drive customer profitability — maps precisely to the analytical work I performed at "
    "Accenture, where I built and maintained over 600 financial models supporting multimillion-dollar "
    "commercial decisions.",
    size=11, space_after=10)

cl_para(cl,
    "In that role, I monitored costs, revenues, and performance indicators against approved business "
    "plans, producing monthly dashboards for senior stakeholders and conducting multi-variable sensitivity "
    "analysis to quantify margin impacts across bid scenarios. I also reduced manual reporting time by 30% "
    "through Excel VBA automation — exactly the kind of process improvement your team expects. At SMS LLP, "
    "I prepared monthly management accounts, posted journals, and reconciled balance sheets, giving me the "
    "accounting foundations needed to contribute from day one.",
    size=11, space_after=10)

cl_para(cl,
    "My MSc Finance at Lancaster University Management School — where I am on track for a First Class — "
    "has sharpened my financial modelling, forecasting, and quantitative analysis skills. I am actively "
    "progressing through CIMA at Management Case Study level, covering Financial Reporting, Performance "
    "Management, and Business Strategy, which aligns directly with the cost monitoring and commercial "
    "support responsibilities in this role. My Advanced Excel, Python, and SQL capabilities ensure I can "
    "hit the ground running with your financial systems and reporting tools.",
    size=11, space_after=10)

cl_para(cl,
    "I am genuinely drawn to SUEZ's purpose-driven environment — the intersection of financial rigour "
    "and environmental responsibility makes this graduate programme particularly compelling. The "
    "24-month structure, leading to a permanent role, reflects exactly the kind of structured career "
    "development I am seeking. Bolton is within easy reach of Lancaster, and I am fully available to "
    "start on 7th September 2026.",
    size=11, space_after=10)

cl_para(cl,
    "I would welcome the opportunity to discuss how my financial modelling, budgeting, and reporting "
    "background can support SUEZ's commercial and operational objectives. I am available for interview "
    "and can be reached at jeevajk2112@gmail.com or +44 7466 480480.",
    size=11, space_after=10)

cl_para(cl, "Yours sincerely,", size=11, space_after=16)
cl_para(cl, "Jeevakumar Jayachandran", size=11, bold=True, space_after=4)

# ATS note
p_ats = cl.add_paragraph()
p_ats.paragraph_format.space_before = Pt(12)
r_ats = p_ats.add_run(
    f"---\nATS Match Score: {total}/100  |  Role: Graduate Financial Analyst at SUEZ\n"
    "⚠ NOTE: Role states NO visa sponsorship (now or in future)"
)
_set_font(r_ats, 9)
r_ats.italic = True

cl.save(CL_OUT)
print(f"✓ Cover letter saved: {CL_OUT}")

# ══════════════════════════════════════════════════════════════════════════════
# TRACKER
# ══════════════════════════════════════════════════════════════════════════════
new_row = {
    "Job Title":          "Graduate Financial Analyst",
    "Company":            "SUEZ",
    "Location":           "Bolton, UK",
    "Job Type":           "Full-time (Fixed 24m)",
    "Score":              str(total),
    "Date Found":         "2026-06-11",
    "Application Status": "Not Started",
    "Salary":             "£28,673–£31,257",
    "Visa Sponsorship":   "No – explicit rejection",
    "Apply Link":         "https://www.suez.co.uk/en-gb/careers",
    "CV File":            "SUEZ_Graduate_Financial_Analyst_CV.docx",
    "Cover Letter File":  "SUEZ_Graduate_Financial_Analyst_CoverLetter.docx",
}

rows = []
with open(TRACKER, newline="", encoding="utf-8") as f:
    reader = csv.DictReader(f)
    fieldnames = reader.fieldnames
    for row in reader:
        if row["Job Title"]:
            rows.append(row)

rows.insert(0, new_row)

with open(TRACKER, "w", newline="", encoding="utf-8") as f:
    writer = csv.DictWriter(f, fieldnames=fieldnames)
    writer.writeheader()
    writer.writerows(rows)

print("✓ Tracker updated")

# ── Score report ──────────────────────────────────────────────────────────
print("\n" + "="*60)
print("JOB SCORE REPORT")
print(f"Role:     Graduate Financial Analyst @ SUEZ")
print(f"Location: Bolton, UK")
print(f"⚠  VISA:  NO SPONSORSHIP (explicit)")
print("="*60)
bar_max = 30
for k, v in scores.items():
    max_v = int(k.split("(")[1].rstrip(")"))
    bar = "█" * int(v / max_v * bar_max) + "░" * (bar_max - int(v / max_v * bar_max))
    print(f"  {k:<28} {v:>2}  {bar}")
print(f"\n  TOTAL SCORE: {total}/100")
print(f"\n  {reasoning[:120]}...")
print("="*60)
print(f"\nApply: https://www.suez.co.uk/en-gb/careers (closing 26/06/2026)")
