"""
Tailored CV + Cover Letter generator for:
  FP&A Graduate @ Alexander Dennis — Skelmersdale, UK
  Hybrid | Full-time | Posted 9 hours ago
"""
import csv, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CV_OUT  = "applications/cvs/Alexander_Dennis_FPandA_Graduate_CV.docx"
CL_OUT  = "applications/coverletters/Alexander_Dennis_FPandA_Graduate_CoverLetter.docx"
TRACKER = "applications/tracker.csv"
os.makedirs("applications/cvs",          exist_ok=True)
os.makedirs("applications/coverletters", exist_ok=True)

scores = {
    "Education Match (20)":    19,
    "Technical Skills (30)":   26,
    "Experience Match (20)":   18,
    "Visa Compatibility (10)":  7,
    "Location Match (10)":      7,
    "Career Alignment (10)":   10,
}
total = sum(scores.values())

FONT = "Calibri"

def _set_font(run, size, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, 10, bold=True, color=(0,70,127))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "00468F")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def two_col(doc, left, right, left_bold=False, right_bold=False,
            left_size=10, right_size=9, left_italic=False, right_italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Twips(int(6.73 * 1440)), 2)
    rl = p.add_run(left)
    _set_font(rl, left_size, bold=left_bold)
    rl.italic = left_italic
    rr = p.add_run("\t" + right)
    _set_font(rr, right_size, bold=right_bold)
    rr.italic = right_italic
    return p

def bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    _set_font(run, size)
    return p

def body_line(doc, text, size=9.5, bold=False, space_after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, size, bold=bold)
    return p

def set_narrow_margins(doc):
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.5)
        sec.left_margin = sec.right_margin = Inches(0.5)

def set_normal_margins(doc):
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1.0)
        sec.left_margin = sec.right_margin = Inches(1.0)

# ══════════════════════════════════════════════════════════════════════════════
# CV
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
set_narrow_margins(doc)

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = name_p.add_run("JEEVAKUMAR JAYACHANDRAN")
_set_font(nr, 16, bold=True, color=(0,70,127))

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(4)
cr = contact_p.add_run(
    "Lancaster, UK  |  jeevajk2112@gmail.com  |  +44 7466 480480  |  linkedin.com/in/jeevakumar-j"
)
_set_font(cr, 9)

# Profile
section_heading(doc, "Professional Profile")
body_line(doc,
    "MSc Finance graduate (First Class expected, Lancaster University) with two years' hands-on experience "
    "in FP&A activities: overhead cost analysis, variance analysis, budgeting, forecasting, and management "
    "reporting. Adept at building financial models in Excel and translating complex data into clear insights "
    "for senior stakeholders. Part-qualified accountant (CIMA Management Case Study level) with strong "
    "commercial awareness developed across manufacturing-adjacent and professional services environments.",
    size=9.5, space_after=2)

# Education
section_heading(doc, "Education")

two_col(doc, "MSc Finance", "Oct 2025 – Aug 2026", left_bold=True, left_size=10)
two_col(doc, "Lancaster University Management School, Lancaster, UK",
        "First Class (Expected)", left_italic=True, right_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Modules: Financial Modelling, Corporate Finance, Investment Analysis, Quantitative Methods, Bloomberg Terminal")
bullet(doc, "Dissertation: Quantitative analysis of ESG factors and equity risk premia across FTSE 350 constituents")

two_col(doc, "CIMA – Management Case Study Level (In Progress)", "2025 – Present", left_bold=True, left_size=10)
two_col(doc, "Chartered Institute of Management Accountants",
        "", left_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Management level: Financial Reporting, Performance Management, Financial Management, Business Strategy")

two_col(doc, "Master of Commerce (M.Com)", "Jun 2019 – May 2021", left_bold=True, left_size=10)
two_col(doc, "University of Madras, India",
        "First Class – 82%", left_italic=True, right_italic=True, left_size=9.5, right_size=9.5)

two_col(doc, "Bachelor of Commerce (B.Com)", "Jun 2016 – May 2019", left_bold=True, left_size=10)
two_col(doc, "University of Madras, India",
        "First Class – 72%", left_italic=True, right_italic=True, left_size=9.5, right_size=9.5)

# Experience
section_heading(doc, "Professional Experience")

two_col(doc, "Pricing Associate (Financial Analyst)", "Jun 2023 – Sep 2025", left_bold=True, left_size=10)
two_col(doc, "Accenture, Bengaluru, India", "", left_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Supported overhead forecasting and cost analysis across 600+ financial models for large outsourcing contracts, producing accurate financial reports used by senior leadership to monitor performance against budget.")
bullet(doc, "Prepared detailed variance analysis reports comparing actuals to forecast, identifying cost drivers and presenting findings to budget holders and senior stakeholders with clear narrative commentary.")
bullet(doc, "Built multi-variable sensitivity models in Excel to assess overhead cost scenarios and margin impacts, supporting commercial decision-making and financial planning cycles.")
bullet(doc, "Partnered with operations and commercial teams to deliver rolling forecasts and annual budgets, translating business assumptions into financial projections aligned with operational plans.")
bullet(doc, "Reduced manual financial reporting time by 30% by automating overhead tracking dashboards using Excel VBA and Power Query, improving data accuracy and turnaround for senior reports.")

two_col(doc, "Finance Assistant", "Feb 2022 – May 2023", left_bold=True, left_size=10)
two_col(doc, "SMS LLP (Mtandt Group), Chennai, India", "", left_italic=True, left_size=9.5, right_size=9.5)
bullet(doc, "Prepared monthly management accounts and variance analysis reports, monitoring actual costs and revenues against budget and providing commentary to support operational decisions.")
bullet(doc, "Conducted detailed inventory and cost analysis that identified a 12% reduction in holding losses, presenting findings with supporting financial models to management.")
bullet(doc, "Assisted in budgeting cycles including cost centre preparation, accruals, prepayments, and balance sheet reconciliations across multiple departments.")
bullet(doc, "Maintained accurate financial records through journal postings and period-end close activities, ensuring data integrity for management reporting.")

# Projects
section_heading(doc, "Academic Projects")

two_col(doc, "Unilever PLC – FCFF Valuation Model", "2026", left_bold=True, left_size=10)
bullet(doc, "Built a three-statement DCF model forecasting free cash flow to firm; benchmarked against analyst consensus using Bloomberg Terminal data.")

two_col(doc, "FTSE 350 – CAPM & Risk Analysis", "2026", left_bold=True, left_size=10)
bullet(doc, "Applied Fama-French three-factor model using R (panel regression) across FTSE 350; analysed systematic risk and equity risk premia.")

# Skills
section_heading(doc, "Skills & Competencies")

two_col(doc, "FP&A:", "Overhead Forecasting, Budgeting, Variance Analysis, Cost Analysis, Management Reporting, KPI Tracking",
        left_bold=True, left_size=9.5, right_size=9.5)
two_col(doc, "Financial Modelling:", "DCF, Sensitivity Analysis, Scenario Modelling, Profitability & Cost Modelling",
        left_bold=True, left_size=9.5, right_size=9.5)
two_col(doc, "Technical Tools:", "Advanced Excel (VBA, Power Query, Pivot Tables), Python (pandas), SQL, R, Power BI, Bloomberg Terminal",
        left_bold=True, left_size=9.5, right_size=9.5)
two_col(doc, "Professional:", "CIMA – Management Case Study Level (in progress) | CMA Intermediate (ICAI)",
        left_bold=True, left_size=9.5, right_size=9.5)

doc.save(CV_OUT)
print(f"✓ CV saved:           {CV_OUT}")

# ══════════════════════════════════════════════════════════════════════════════
# COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
cl = Document()
set_normal_margins(cl)

def cl_para(doc, text, size=11, bold=False, space_after=8, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _set_font(r, size, bold=bold)
    return p

cl_para(cl, "Jeevakumar Jayachandran", size=13, bold=True, space_after=2)
cl_para(cl, "Lancaster, UK  |  jeevajk2112@gmail.com  |  +44 7466 480480  |  linkedin.com/in/jeevakumar-j",
        size=10, space_after=12)
cl_para(cl, "2026-06-11", size=11, space_after=8)
cl_para(cl, "Hiring Manager\nAlexander Dennis\nSkelmersdale, UK", size=11, space_after=12)
cl_para(cl, "Re: Application for FP&A Graduate", size=11, bold=True, space_after=12)
cl_para(cl, "Dear Hiring Manager,", size=11, space_after=12)

cl_para(cl,
    "Alexander Dennis's mission to lead sustainable mass transportation — and the financial rigour "
    "required to support that at scale — is exactly the kind of environment I'm looking to grow in. "
    "The FP&A Graduate role, with its focus on overhead forecasting, cost and inventory analysis, and "
    "reporting to senior leadership, maps directly to the work I've been doing professionally for the "
    "past two years.",
    size=11, space_after=10)

cl_para(cl,
    "At Accenture, I maintained over 600 financial models for large commercial contracts, built overhead "
    "cost forecasts, and produced variance analysis reports comparing actuals to plan for senior "
    "stakeholders. I worked closely with budget holders across finance, operations, and commercial "
    "teams — the same cross-functional partnering this role requires. I also reduced manual reporting "
    "time by 30% through Excel VBA automation, which reflects the kind of process improvement mindset "
    "I bring to any finance team. At SMS LLP, I prepared monthly management accounts, conducted "
    "inventory and cost analysis, and supported full budgeting cycles.",
    size=11, space_after=10)

cl_para(cl,
    "My MSc Finance at Lancaster University Management School (on track for a First Class) has "
    "strengthened my financial modelling and quantitative skills, and I am currently progressing "
    "through CIMA at Management Case Study level — covering Performance Management, Financial "
    "Reporting, and Business Strategy. This part-qualification, combined with hands-on FP&A "
    "experience, means I can contribute meaningfully from the outset rather than needing an "
    "extended ramp-up period.",
    size=11, space_after=10)

cl_para(cl,
    "I'm drawn to Alexander Dennis specifically because the financial role directly supports "
    "something tangible — buses that reduce emissions and keep communities connected. That sense "
    "of purpose, alongside the structured development on offer, makes this a programme I would "
    "commit to fully.",
    size=11, space_after=10)

cl_para(cl,
    "I would welcome the opportunity to discuss my application further and am available for "
    "interview at your convenience. Please feel free to contact me at jeevajk2112@gmail.com "
    "or +44 7466 480480.",
    size=11, space_after=10)

cl_para(cl, "Yours sincerely,", size=11, space_after=16)
cl_para(cl, "Jeevakumar Jayachandran", size=11, bold=True, space_after=4)

p_ats = cl.add_paragraph()
p_ats.paragraph_format.space_before = Pt(12)
r_ats = p_ats.add_run(
    f"---\nATS Match Score: {total}/100  |  Role: FP&A Graduate at Alexander Dennis"
)
_set_font(r_ats, 9)
r_ats.italic = True

cl.save(CL_OUT)
print(f"✓ Cover letter saved: {CL_OUT}")

# ══════════════════════════════════════════════════════════════════════════════
# TRACKER
# ══════════════════════════════════════════════════════════════════════════════
new_row = {
    "Job Title":          "FP&A Graduate",
    "Company":            "Alexander Dennis",
    "Location":           "Skelmersdale, UK",
    "Job Type":           "Hybrid",
    "Score":              str(total),
    "Date Found":         "2026-06-11",
    "Application Status": "Not Started",
    "Salary":             "Competitive",
    "Visa Sponsorship":   "Check listing",
    "Apply Link":         "https://uk.linkedin.com/jobs/view/fpa-graduate-at-alexander-dennis-limited",
    "CV File":            "Alexander_Dennis_FPandA_Graduate_CV.docx",
    "Cover Letter File":  "Alexander_Dennis_FPandA_Graduate_CoverLetter.docx",
}

rows = []
with open(TRACKER, newline="", encoding="utf-8") as f:
    reader = csv.DictReader(f)
    fieldnames = reader.fieldnames
    for row in reader:
        if row["Job Title"]:
            rows.append(row)

rows.insert(0, new_row)

with open(TRACKER, "w", newline="", encoding="utf-8") as f:
    writer = csv.DictWriter(f, fieldnames=fieldnames)
    writer.writeheader()
    writer.writerows(rows)

print("✓ Tracker updated")

print("\n" + "="*60)
print("JOB SCORE REPORT")
print("Role:     FP&A Graduate @ Alexander Dennis")
print("Location: Skelmersdale, UK (Hybrid)")
print("="*60)
for k, v in scores.items():
    max_v = int(k.split("(")[1].rstrip(")"))
    bar = "█" * int(v / max_v * 30) + "░" * (30 - int(v / max_v * 30))
    print(f"  {k:<28} {v:>2}  {bar}")
print(f"\n  TOTAL SCORE: {total}/100")
print("="*60)
